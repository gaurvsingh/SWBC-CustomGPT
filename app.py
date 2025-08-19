# # app.py  — SWBC RFP Assistant (Frontend + Backend)
# # End-to-end: Ingest (reference boost) → Drafting → QA (validation/grounding/privacy) → DOCX export

# import os, io, re, uuid, json
# from dataclasses import dataclass, asdict
# from typing import Any, Dict, List, Optional, Tuple

# import numpy as np
# from flask import Flask, request, jsonify, render_template, send_file
# from flask_cors import CORS
# from dotenv import load_dotenv
# from openai import OpenAI
# from PyPDF2 import PdfReader

# try:
#     from docx import Document
# except Exception:
#     Document = None  # DOCX export disabled if python-docx isn't installed

# # =========================
# # ENV + CLIENT
# # =========================
# load_dotenv()

# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# if not OPENAI_API_KEY:
#     raise RuntimeError("OPENAI_API_KEY is required in environment")

# client = OpenAI(api_key=OPENAI_API_KEY)

# DATA_DIR = os.getenv("DATA_DIR", "./data")
# EXPORT_DIR = os.getenv("EXPORT_DIR", "./exports")
# os.makedirs(DATA_DIR, exist_ok=True)
# os.makedirs(EXPORT_DIR, exist_ok=True)

# # Optional auth for API calls (UI routes allowed without)
# ACTION_API_KEY = os.getenv("API_KEY")

# # Reference docs: boosted in retrieval (helps tone/structure)
# REFERENCE_DOC_IDS = {
#     x.strip() for x in os.getenv("REFERENCE_DOC_IDS", "RFQ_C2.pdf").split(",") if x.strip()
# }
# REFERENCE_BOOST = float(os.getenv("REFERENCE_BOOST", "1.25"))

# VECTOR_BACKEND = os.getenv("VECTOR_BACKEND", "local").lower()

# # Models
# EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")
# DRAFT_MODEL = os.getenv("DRAFT_MODEL", "gpt-4o-mini")
# EXTRACT_MODEL = os.getenv("EXTRACT_MODEL", "gpt-4o-mini")

# # =========================
# # UTILS
# # =========================
# def api_guard_ok():
#     if ACTION_API_KEY and request.headers.get("X-API-KEY") != ACTION_API_KEY:
#         return False
#     return True

# def l2norm(x: np.ndarray) -> np.ndarray:
#     n = np.linalg.norm(x)
#     return x / (n + 1e-12)

# def cosine(a: np.ndarray, b: np.ndarray) -> float:
#     return float(np.dot(l2norm(a), l2norm(b)))

# def embed_text(text: str) -> List[float]:
#     # strip images/nulls & clamp
#     text = (text or "").replace("\x00", " ")
#     text = re.sub(r"<img[^>]*>", "", text)
#     text = text[:8000]
#     resp = client.embeddings.create(model=EMBED_MODEL, input=text)
#     return resp.data[0].embedding

# # =========================
# # LOCAL VECTOR STORE
# # =========================
# CHUNK_FILE = os.path.join(DATA_DIR, "chunks.jsonl")

# @dataclass
# class Chunk:
#     chunk_id: str
#     doc_id: str
#     page: int
#     offset: Tuple[int, int]
#     text: str
#     meta: Dict[str, Any]
#     vector: Optional[List[float]] = None
#     score: Optional[float] = None

# class VectorStore:
#     def upsert(self, chunks: List[Chunk]) -> None: ...
#     def search(self, query: str, k: int = 6, filters: Optional[Dict[str, Any]] = None) -> List[Chunk]: ...
#     def has_doc(self, doc_id: str) -> bool: ...
#     def get_doc_page_chunks(self, doc_id: str, page: Optional[int], limit: int = 20) -> List[Chunk]:
#         return []

# class LocalJSONStore(VectorStore):
#     def upsert(self, chunks: List[Chunk]) -> None:
#         # Append JSONL; embed per chunk (keeps memory flat)
#         with open(CHUNK_FILE, "a", encoding="utf-8") as f:
#             for c in chunks:
#                 if c.vector is None:
#                     c.vector = embed_text(c.text)
#                 f.write(json.dumps(asdict(c)) + "\n")

#     def _iter_chunks(self):
#         if not os.path.exists(CHUNK_FILE):
#             return
#         with open(CHUNK_FILE, "r", encoding="utf-8") as f:
#             for line in f:
#                 line = line.strip()
#                 if not line:
#                     continue
#                 yield Chunk(**json.loads(line))

#     def has_doc(self, doc_id: str) -> bool:
#         for c in self._iter_chunks():
#             if c.doc_id == doc_id:
#                 return True
#         return False

#     def search(self, query: str, k: int = 6, filters: Optional[Dict[str, Any]] = None) -> List[Chunk]:
#         qvec = np.array(embed_text(query), dtype=np.float32)
#         prefer_reference = (filters or {}).get("reference_first", False)
#         results: List[Chunk] = []
#         for c in self._iter_chunks():
#             cvec = np.array(c.vector, dtype=np.float32)
#             s = cosine(qvec, cvec)
#             if prefer_reference or c.doc_id in REFERENCE_DOC_IDS or c.meta.get("reference") is True:
#                 s *= REFERENCE_BOOST
#             c.score = s
#             results.append(c)
#         results.sort(key=lambda x: x.score or 0.0, reverse=True)
#         return results[:max(1, min(k, 50))]

#     def get_doc_page_chunks(self, doc_id: str, page: Optional[int], limit: int = 20) -> List[Chunk]:
#         out: List[Chunk] = []
#         for c in self._iter_chunks():
#             if c.doc_id != doc_id:
#                 continue
#             if page is not None and c.page != page:
#                 continue
#             out.append(c)
#             if len(out) >= limit:
#                 break
#         return out

# # Backends (stubs to make switching easy later)
# class QdrantStore(VectorStore): ...
# class PgVectorStore(VectorStore): ...
# class AtlasVectorStore(VectorStore): ...

# def make_store() -> VectorStore:
#     if VECTOR_BACKEND == "qdrant": return QdrantStore()
#     if VECTOR_BACKEND == "pgvector": return PgVectorStore()
#     if VECTOR_BACKEND == "atlas": return AtlasVectorStore()
#     return LocalJSONStore()

# store: VectorStore = make_store()

# # =========================
# # INGESTION (memory-safe)
# # =========================
# CHUNK_LEN = 900
# CHUNK_OVERLAP = 120

# def iter_chunks(txt: str, chunk_len: int = CHUNK_LEN, overlap: int = CHUNK_OVERLAP):
#     n = len(txt)
#     i = 0
#     while i < n:
#         j = min(n, i + chunk_len)
#         yield (i, j, txt[i:j])
#         if j >= n:
#             break
#         i = max(0, j - overlap)

# def clean_page_text(t: str) -> str:
#     t = (t or "").replace("\x00", " ")
#     if len(t) > 600_000:  # clamp pathological PDF pages
#         t = t[:600_000]
#     return t

# def pdf_to_pages(path: str) -> List[Tuple[int, str]]:
#     reader = PdfReader(path)
#     pages = []
#     for p, page in enumerate(reader.pages):
#         try:
#             t = page.extract_text() or ""
#         except Exception:
#             t = ""
#         pages.append((p, t))
#     return pages

# def docx_to_text(path: str) -> str:
#     if not Document:
#         raise RuntimeError("python-docx is required for DOCX ingestion. pip install python-docx")
#     doc = Document(path)
#     return "\n".join([p.text for p in doc.paragraphs])

# def ingest_file(filepath: str, reference: bool = False) -> int:
#     """
#     Streams chunks to the vector store in small batches (low RAM).
#     """
#     doc_id = os.path.basename(filepath)
#     total_chunks = 0

#     def _upsert_batch(batch: List[Chunk]):
#         nonlocal total_chunks
#         if not batch: return
#         store.upsert(batch)
#         total_chunks += len(batch)
#         batch.clear()

#     batch: List[Chunk] = []
#     BATCH_SIZE = 16
#     ref_flag = reference or (doc_id in REFERENCE_DOC_IDS)

#     if filepath.lower().endswith(".pdf"):
#         for page_idx, text in pdf_to_pages(filepath):
#             text = clean_page_text(text)
#             if not text.strip():
#                 continue
#             for start, end, span in iter_chunks(text):
#                 batch.append(Chunk(
#                     chunk_id=str(uuid.uuid4()),
#                     doc_id=doc_id,
#                     page=page_idx,
#                     offset=(start, end),
#                     text=span,
#                     meta={"reference": ref_flag}
#                 ))
#                 if len(batch) >= BATCH_SIZE:
#                     _upsert_batch(batch)

#     elif filepath.lower().endswith(".docx"):
#         text = clean_page_text(docx_to_text(filepath))
#         for start, end, span in iter_chunks(text):
#             batch.append(Chunk(
#                 chunk_id=str(uuid.uuid4()),
#                 doc_id=doc_id,
#                 page=0,
#                 offset=(start, end),
#                 text=span,
#                 meta={"reference": ref_flag}
#             ))
#             if len(batch) >= BATCH_SIZE:
#                 _upsert_batch(batch)

#     else:
#         with open(filepath, "r", encoding="utf-8") as f:
#             text = clean_page_text(f.read())
#         for start, end, span in iter_chunks(text):
#             batch.append(Chunk(
#                 chunk_id=str(uuid.uuid4()),
#                 doc_id=doc_id,
#                 page=0,
#                 offset=(start, end),
#                 text=span,
#                 meta={"reference": ref_flag}
#             ))
#             if len(batch) >= BATCH_SIZE:
#                 _upsert_batch(batch)

#     _upsert_batch(batch)  # flush
#     return total_chunks

# # =========================
# # QA HELPERS (privacy + basic grounding)
# # =========================
# EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
# PHONE_RE = re.compile(r"\b(?:\+?\d{1,2}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}\b")
# SSN_RE   = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
# CC_RE    = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
# APIKEY_RE= re.compile(r"\b(?:sk-[A-Za-z0-9]{20,}|AKIA[0-9A-Z]{16})\b")

# def luhn_ok(s: str) -> bool:
#     digits = [int(d) for d in re.sub(r"\D", "", s)]
#     if len(digits) < 13: return False
#     checksum = 0
#     parity = len(digits) % 2
#     for i, d in enumerate(digits):
#         if i % 2 == parity:
#             d *= 2
#             if d > 9: d -= 9
#         checksum += d
#     return (checksum % 10) == 0

# def redact_answer(ans: str):
#     redactions = []
#     def _re(regex, token):
#         nonlocal ans
#         for m in list(regex.finditer(ans)):
#             redactions.append({"type": token.strip("<>"), "span": [m.start(), m.end()]})
#         ans = regex.sub(token, ans)
#     _re(EMAIL_RE, "<REDACTED_EMAIL>")
#     _re(PHONE_RE, "<REDACTED_PHONE>")
#     _re(SSN_RE, "<REDACTED_SSN>")
#     for m in list(CC_RE.finditer(ans)):
#         if luhn_ok(m.group(0)):
#             redactions.append({"type":"CREDIT_CARD","span":[m.start(),m.end()]})
#             ans = ans[:m.start()] + "<REDACTED_CC>" + ans[m.end():]
#     _re(APIKEY_RE, "<REDACTED_SECRET>")
#     return ans, redactions

# def check_grounding(answer: str, citations: List[Dict[str, Any]]) -> Tuple[bool, List[str]]:
#     if not answer.strip():
#         return False, ["Empty answer."]
#     if not citations:
#         return False, ["No citations provided."]
#     # Simple lexical overlap with cited chunks from same doc/page
#     ans_tokens = re.findall(r"[A-Za-z0-9]+", answer.lower())
#     if not ans_tokens:
#         return False, ["Unusable answer tokens."]
#     ans_grams = set(" ".join(ans_tokens[i:i+4]) for i in range(max(1, len(ans_tokens) - 3)))
#     overlap = 0
#     for cite in citations:
#         doc_id = cite.get("doc_id")
#         page = cite.get("page")
#         # Prefer direct access to doc/page chunks:
#         if hasattr(store, "get_doc_page_chunks"):
#             candidates = store.get_doc_page_chunks(doc_id, page, limit=20)
#         else:
#             candidates = store.search(f"{doc_id} page {page}", k=20)
#         for ch in candidates:
#             c_tokens = re.findall(r"[A-Za-z0-9]+", (ch.text or "").lower())
#             c_grams = set(" ".join(c_tokens[i:i+4]) for i in range(max(1, len(c_tokens) - 3)))
#             if ans_grams & c_grams:
#                 overlap += 1
#                 break
#     if overlap == 0:
#         return False, ["Answer appears ungrounded in cited snippets."]
#     return True, []

# # =========================
# # DRAFTING (LLM)
# # =========================
# DRAFT_SYSTEM = (
#     "You are the Drafting Agent. Use ONLY the provided context snippets to answer. "
#     "Do not invent facts. Keep tone professional. "
#     "Return ONLY JSON: {\"answer\":\"...\",\"status\":\"ok|needs_more_context\"}"
# )

# def extract_questions_llm(text: str) -> List[str]:
#     prompt = f"""
# Extract EVERY question from this RFP text. Keep multi-part items together.
# Output ONLY a JSON array of strings.

# Text:
# {text[:120000]}
# """
#     resp = client.chat.completions.create(
#         model=EXTRACT_MODEL,
#         temperature=0,
#         messages=[
#             {"role":"system","content":"You extract RFP questions precisely and output valid JSON."},
#             {"role":"user","content":prompt}
#         ]
#     )
#     raw = (resp.choices[0].message.content or "").strip().strip("`").strip()
#     try:
#         arr = json.loads(raw)
#         return [q.strip() for q in arr if isinstance(q, str) and q.strip()]
#     except Exception:
#         # Fallback: naive tokenization on '?'
#         parts = [p.strip()+"?" for p in text.split("?") if p.strip()]
#         return parts[:25]

# def draft_answer_llm(question: str, ctx_chunks: List[Dict[str,Any]]) -> Dict[str, Any]:
#     context = "\n\n".join([f"[{i+1}] {c['text']}" for i,c in enumerate(ctx_chunks[:6])])
#     prompt = f"Question: {question}\n\nContext:\n{context}"
#     resp = client.chat.completions.create(
#         model=DRAFT_MODEL,
#         temperature=0.1,
#         messages=[{"role":"system","content":DRAFT_SYSTEM},
#                   {"role":"user","content":prompt}]
#     )
#     raw = (resp.choices[0].message.content or "").strip().strip("`").strip()
#     try:
#         return json.loads(raw)
#     except Exception:
#         return {"answer": raw, "status": "ok"}

# def pick_citations(ctx_chunks: List[Dict[str,Any]], n: int = 2) -> List[Dict[str,Any]]:
#     cites = []
#     for c in ctx_chunks[:max(1, n)]:
#         cites.append({
#             "doc_id": c.get("doc_id"),
#             "page": c.get("page"),
#             "offset": c.get("offset")
#         })
#     return cites

# def qa_answer(answer: str, citations: List[Dict[str,Any]]) -> Dict[str,Any]:
#     status = "pass"
#     reasons: List[str] = []

#     ok_g, g_reasons = check_grounding(answer, citations)
#     if not ok_g:
#         status = "fail"
#         reasons.extend(g_reasons)

#     redacted, redactions = redact_answer(answer)
#     risk = 0
#     if status == "fail": risk += 60
#     if redactions: risk += 20

#     fix = None
#     if status == "fail":
#         fix = ("Provide a concise answer grounded strictly in the cited snippets; "
#                "quote specific phrases and include numeric details when present.")

#     return {
#         "status": status,
#         "reasons": reasons,
#         "fix": fix,
#         "redactions": redactions,
#         "controls": [],
#         "risk": risk,
#         "answer_effective": redacted
#     }

# # =========================
# # EXPORT
# # =========================
# def export_docx(rfp_id: str, qaed: List[Dict[str,Any]]) -> str:
#     if not Document:
#         raise RuntimeError("python-docx not installed. Run: pip install python-docx")
#     doc = Document()
#     doc.add_heading(f"RFP Response – {rfp_id}", level=1)
#     for i, item in enumerate(qaed, start=1):
#         q = item["question"]
#         a = item["qa"]["answer_effective"] or item["draft"]["answer"] or ""
#         doc.add_heading(f"Q{i}. {q}", level=2)
#         for para in (a or "").split("\n\n"):
#             doc.add_paragraph(para)
#     out_path = os.path.join(EXPORT_DIR, f"{rfp_id}.docx")
#     doc.save(out_path)
#     return out_path

# # =========================
# # FLASK APP
# # =========================
# app = Flask(__name__, static_folder="static", template_folder="templates")
# CORS(app)
# app.config["MAX_CONTENT_LENGTH"] = 40 * 1024 * 1024  # 40 MB uploads

# @app.before_request
# def _guard():
#     # Allow UI & health & ingest (for quick testing) without API key
#     if request.path in ("/", "/health", "/ingest/upload") or request.path.startswith("/package/"):
#         return
#     if not api_guard_ok():
#         return jsonify({"error": "unauthorized"}), 401

# @app.get("/health")
# def health():
#     return jsonify({"ok": True, "vector_backend": VECTOR_BACKEND})

# @app.get("/")
# def index():
#     return render_template("index.html")

# @app.post("/ingest/upload")
# def ingest_upload():
#     f = request.files.get("file")
#     if not f:
#         return jsonify({"error": "No file"}), 400
#     filename = f.filename or f"upload-{uuid.uuid4()}"
#     path = os.path.join(DATA_DIR, filename)
#     f.save(path)

#     # reference flag via query or form
#     reference = (
#         request.args.get("reference", "false").lower() in ("1", "true", "yes") or
#         request.form.get("reference", "false").lower() in ("1", "true", "yes")
#     )
#     try:
#         count = ingest_file(path, reference=reference)
#     except Exception as e:
#         return jsonify({"error": f"Ingestion failed: {e}"}), 500
#     return jsonify({"doc_id": filename, "chunks": count, "reference": reference})

# @app.post("/rfp/complete")
# def rfp_complete():
#     """
#     Upload an RFP → extract questions → for each:
#       - retrieve context (reference-first)
#       - draft answer (LLM)
#       - QA (validation/grounding/privacy)
#     → export DOCX + return preview JSON
#     """
#     f = request.files.get("rfp")
#     if not f:
#         return jsonify({"error": "No RFP uploaded"}), 400

#     limit = int(request.form.get("limit", "20"))
#     kctx  = int(request.form.get("k", "6"))

#     # Save and read text
#     tmpname = f"rfp-{uuid.uuid4()}{os.path.splitext(f.filename or '')[1]}"
#     rfp_path = os.path.join(DATA_DIR, tmpname)
#     f.save(rfp_path)

#     if rfp_path.lower().endswith(".pdf"):
#         pages = pdf_to_pages(rfp_path)
#         rfp_text = "\n\n".join([t for _, t in pages])
#     elif rfp_path.lower().endswith(".docx"):
#         rfp_text = docx_to_text(rfp_path)
#     else:
#         rfp_text = open(rfp_path, "r", encoding="utf-8").read()

#     # 1) Extract questions
#     questions = extract_questions_llm(rfp_text)[:limit]
#     results: List[Dict[str,Any]] = []

#     # 2) Draft + QA per question
#     for q in questions:
#         ctx = store.search(q, k=kctx, filters={"reference_first": True})
#         ctx_payload = [{
#             "doc_id": c.doc_id,
#             "page": c.page,
#             "offset": list(c.offset),
#             "text": c.text
#         } for c in ctx]

#         draft = draft_answer_llm(q, ctx_payload)
#         ans = (draft.get("answer") or "").strip()
#         status = draft.get("status", "ok")

#         citations = pick_citations(ctx_payload, n=2)
#         qa = qa_answer(ans, citations)

#         # Optional single-pass auto-fix if fail
#         if qa["status"] == "fail" and qa.get("fix"):
#             qa2 = qa_answer(qa["fix"], citations)
#             if qa2["status"] == "pass":
#                 qa = qa2

#         results.append({
#             "question": q,
#             "citations": citations,
#             "draft": {"status": status, "answer": ans},
#             "qa": qa
#         })

#     # 3) Export DOCX
#     rfp_id = os.path.splitext(os.path.basename(rfp_path))[0]
#     try:
#         out_path = export_docx(rfp_id, results)
#     except Exception as e:
#         # Still return answers even if export fails
#         return jsonify({
#             "rfp_id": rfp_id,
#             "count_questions": len(questions),
#             "answers": results,
#             "download_uri": None,
#             "export_error": str(e)
#         })

#     return jsonify({
#         "rfp_id": rfp_id,
#         "count_questions": len(questions),
#         "answers": results,
#         "download_uri": f"/package/download/{os.path.basename(out_path)}"
#     })

# @app.get("/package/download/<fname>")
# def package_download(fname):
#     path = os.path.join(EXPORT_DIR, fname)
#     if not os.path.exists(path):
#         return jsonify({"error": "not found"}), 404
#     return send_file(path, as_attachment=True)

# if __name__ == "__main__":
#     port = int(os.getenv("PORT", "5000"))
#     debug_mode = os.getenv("FLASK_DEBUG", "1") in ("1", "true", "True")
#     app.run(host="0.0.0.0", port=port, debug=debug_mode)

#v2
# app.py  — SWBC RFP Assistant (Frontend + Backend)
# End-to-end: Ingest (reference boost) → Drafting → QA (validation/grounding/privacy/compliance) → DOCX export

import os, io, re, uuid, json, unicodedata
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader

try:
    from docx import Document
except Exception:
    Document = None  # DOCX export disabled if python-docx isn't installed

# =========================
# ENV + CLIENT
# =========================
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY is required in environment")

client = OpenAI(api_key=OPENAI_API_KEY)

DATA_DIR = os.getenv("DATA_DIR", "./data")
EXPORT_DIR = os.getenv("EXPORT_DIR", "./exports")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

# Optional auth for API calls (UI routes allowed without)
ACTION_API_KEY = os.getenv("API_KEY")

# Reference docs: boosted in retrieval (helps tone/structure)
REFERENCE_DOC_IDS = {
    x.strip() for x in os.getenv("REFERENCE_DOC_IDS", "RFQ_C2.pdf").split(",") if x.strip()
}
REFERENCE_BOOST = float(os.getenv("REFERENCE_BOOST", "1.25"))

VECTOR_BACKEND = os.getenv("VECTOR_BACKEND", "local").lower()

# Models
EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")
DRAFT_MODEL = os.getenv("DRAFT_MODEL", "gpt-4o-mini")
EXTRACT_MODEL = os.getenv("EXTRACT_MODEL", "gpt-4o-mini")

# Optional compliance rules from env (JSON array)
# Example:
# COMPLIANCE_RULES_JSON='[{"when_any":["encryption"],"must_include":["AES-256","TLS 1.2+"],"control":"ISO27001 A.10"}]'
try:
    COMPLIANCE_RULES = json.loads(os.getenv("COMPLIANCE_RULES_JSON", "[]"))
except Exception:
    COMPLIANCE_RULES = []

# =========================
# UTILS
# =========================
def api_guard_ok():
    if ACTION_API_KEY and request.headers.get("X-API-KEY") != ACTION_API_KEY:
        return False
    return True

def l2norm(x: np.ndarray) -> np.ndarray:
    n = np.linalg.norm(x)
    return x / (n + 1e-12)

def cosine(a: np.ndarray, b: np.ndarray) -> float:
    return float(np.dot(l2norm(a), l2norm(b)))

def embed_text(text: str) -> List[float]:
    # strip images/nulls & clamp
    text = (text or "").replace("\x00", " ")
    text = re.sub(r"<img[^>]*>", "", text)
    text = text[:8000]
    resp = client.embeddings.create(model=EMBED_MODEL, input=text)
    return resp.data[0].embedding

def coerce_str(x) -> str:
    """Make any model output string-safe."""
    if x is None:
        return ""
    if isinstance(x, str):
        return x
    if isinstance(x, (list, tuple)):
        parts = []
        for item in x:
            if isinstance(item, dict):
                parts.append(json.dumps(item, ensure_ascii=False))
            else:
                parts.append(coerce_str(item))
        return "\n".join([p for p in parts if p])
    if isinstance(x, dict):
        for key in ("text", "content", "value"):
            v = x.get(key)
            if isinstance(v, str):
                return v
        return json.dumps(x, ensure_ascii=False)
    return str(x)

def normalize_text(s: str) -> str:
    """Ensure consistent characters & whitespace (NFKC, straight quotes, single spacing)."""
    if not s:
        return ""
    # Unicode normalization
    s = unicodedata.normalize("NFKC", s)
    # Curly quotes/dashes → straight
    s = s.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    s = s.replace("–", "-").replace("—", "-")
    # Bullets → "- "
    s = re.sub(r"^[•·]\s*", "- ", s, flags=re.MULTILINE)
    # Collapse excessive blank lines and spaces
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    # Strip weird non-printables
    s = s.replace("\x00", " ")
    return s.strip()

# =========================
# LOCAL VECTOR STORE
# =========================
CHUNK_FILE = os.path.join(DATA_DIR, "chunks.jsonl")

@dataclass
class Chunk:
    chunk_id: str
    doc_id: str
    page: int
    offset: Tuple[int, int]
    text: str
    meta: Dict[str, Any]
    vector: Optional[List[float]] = None
    score: Optional[float] = None

class VectorStore:
    def upsert(self, chunks: List[Chunk]) -> None: ...
    def search(self, query: str, k: int = 6, filters: Optional[Dict[str, Any]] = None) -> List[Chunk]: ...
    def has_doc(self, doc_id: str) -> bool: ...
    def get_doc_page_chunks(self, doc_id: str, page: Optional[int], limit: int = 20) -> List[Chunk]:
        return []

class LocalJSONStore(VectorStore):
    def upsert(self, chunks: List[Chunk]) -> None:
        # Append JSONL; embed per chunk (keeps memory flat)
        with open(CHUNK_FILE, "a", encoding="utf-8") as f:
            for c in chunks:
                if c.vector is None:
                    c.vector = embed_text(c.text)
                f.write(json.dumps(asdict(c)) + "\n")

    def _iter_chunks(self):
        if not os.path.exists(CHUNK_FILE):
            return
        with open(CHUNK_FILE, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                yield Chunk(**json.loads(line))

    def has_doc(self, doc_id: str) -> bool:
        for c in self._iter_chunks():
            if c.doc_id == doc_id:
                return True
        return False

    def search(self, query: str, k: int = 6, filters: Optional[Dict[str, Any]] = None) -> List[Chunk]:
        qvec = np.array(embed_text(query), dtype=np.float32)
        prefer_reference = (filters or {}).get("reference_first", False)
        results: List[Chunk] = []
        for c in self._iter_chunks():
            cvec = np.array(c.vector, dtype=np.float32)
            s = cosine(qvec, cvec)
            if prefer_reference or c.doc_id in REFERENCE_DOC_IDS or c.meta.get("reference") is True:
                s *= REFERENCE_BOOST
            c.score = s
            results.append(c)
        results.sort(key=lambda x: x.score or 0.0, reverse=True)
        return results[:max(1, min(k, 50))]

    def get_doc_page_chunks(self, doc_id: str, page: Optional[int], limit: int = 20) -> List[Chunk]:
        out: List[Chunk] = []
        for c in self._iter_chunks():
            if c.doc_id != doc_id:
                continue
            if page is not None and c.page != page:
                continue
            out.append(c)
            if len(out) >= limit:
                break
        return out

# Backends (stubs to make switching easy later)
class QdrantStore(VectorStore): ...
class PgVectorStore(VectorStore): ...
class AtlasVectorStore(VectorStore): ...

def make_store() -> VectorStore:
    if VECTOR_BACKEND == "qdrant": return QdrantStore()
    if VECTOR_BACKEND == "pgvector": return PgVectorStore()
    if VECTOR_BACKEND == "atlas": return AtlasVectorStore()
    return LocalJSONStore()

store: VectorStore = make_store()

# =========================
# INGESTION (memory-safe)
# =========================
CHUNK_LEN = 900
CHUNK_OVERLAP = 120

def iter_chunks(txt: str, chunk_len: int = CHUNK_LEN, overlap: int = CHUNK_OVERLAP):
    n = len(txt)
    i = 0
    while i < n:
        j = min(n, i + chunk_len)
        yield (i, j, txt[i:j])
        if j >= n:
            break
        i = max(0, j - overlap)

def clean_page_text(t: str) -> str:
    t = (t or "").replace("\x00", " ")
    if len(t) > 600_000:  # clamp pathological PDF pages
        t = t[:600_000]
    return t

def pdf_to_pages(path: str) -> List[Tuple[int, str]]:
    reader = PdfReader(path)
    pages = []
    for p, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        pages.append((p, t))
    return pages

def docx_to_text(path: str) -> str:
    if not Document:
        raise RuntimeError("python-docx is required for DOCX ingestion. pip install python-docx")
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def ingest_file(filepath: str, reference: bool = False) -> int:
    """
    Streams chunks to the vector store in small batches (low RAM).
    """
    doc_id = os.path.basename(filepath)
    total_chunks = 0

    def _upsert_batch(batch: List[Chunk]):
        nonlocal total_chunks
        if not batch: return
        store.upsert(batch)
        total_chunks += len(batch)
        batch.clear()

    batch: List[Chunk] = []
    BATCH_SIZE = 16
    ref_flag = reference or (doc_id in REFERENCE_DOC_IDS)

    if filepath.lower().endswith(".pdf"):
        for page_idx, text in pdf_to_pages(filepath):
            text = clean_page_text(text)
            if not text.strip():
                continue
            for start, end, span in iter_chunks(text):
                batch.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    page=page_idx,
                    offset=(start, end),
                    text=span,
                    meta={"reference": ref_flag}
                ))
                if len(batch) >= BATCH_SIZE:
                    _upsert_batch(batch)

    elif filepath.lower().endswith(".docx"):
        text = clean_page_text(docx_to_text(filepath))
        for start, end, span in iter_chunks(text):
            batch.append(Chunk(
                chunk_id=str(uuid.uuid4()),
                doc_id=doc_id,
                page=0,
                offset=(start, end),
                text=span,
                meta={"reference": ref_flag}
            ))
            if len(batch) >= BATCH_SIZE:
                _upsert_batch(batch)

    else:
        with open(filepath, "r", encoding="utf-8") as f:
            text = clean_page_text(f.read())
        for start, end, span in iter_chunks(text):
            batch.append(Chunk(
                chunk_id=str(uuid.uuid4()),
                doc_id=doc_id,
                page=0,
                offset=(start, end),
                text=span,
                meta={"reference": ref_flag}
            ))
            if len(batch) >= BATCH_SIZE:
                _upsert_batch(batch)

    _upsert_batch(batch)  # flush
    return total_chunks

# =========================
# QA HELPERS (privacy + grounding + validation + compliance)
# =========================
EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"\b(?:\+?\d{1,2}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}\b")
SSN_RE   = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
CC_RE    = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
APIKEY_RE= re.compile(r"\b(?:sk-[A-Za-z0-9]{20,}|AKIA[0-9A-Z]{16})\b")

def luhn_ok(s: str) -> bool:
    digits = [int(d) for d in re.sub(r"\D", "", s)]
    if len(digits) < 13: return False
    checksum = 0
    parity = len(digits) % 2
    for i, d in enumerate(digits):
        if i % 2 == parity:
            d *= 2
            if d > 9: d -= 9
        checksum += d
    return (checksum % 10) == 0

def redact_answer(ans: str):
    redactions = []
    def _re(regex, token):
        nonlocal ans
        for m in list(regex.finditer(ans)):
            redactions.append({"type": token.strip("<>"), "span": [m.start(), m.end()]})
        ans = regex.sub(token, ans)
    _re(EMAIL_RE, "<REDACTED_EMAIL>")
    _re(PHONE_RE, "<REDACTED_PHONE>")
    _re(SSN_RE, "<REDACTED_SSN>")
    for m in list(CC_RE.finditer(ans)):
        if luhn_ok(m.group(0)):
            redactions.append({"type":"CREDIT_CARD","span":[m.start(),m.end()]})
            ans = ans[:m.start()] + "<REDACTED_CC>" + ans[m.end():]
    _re(APIKEY_RE, "<REDACTED_SECRET>")
    return ans, redactions

def check_grounding(answer: str, citations: List[Dict[str, Any]]) -> Tuple[bool, List[str]]:
    if not answer.strip():
        return False, ["Empty answer."]
    if not citations:
        return False, ["No citations provided."]
    # Simple lexical overlap with cited chunks from same doc/page
    ans_tokens = re.findall(r"[A-Za-z0-9]+", answer.lower())
    if not ans_tokens:
        return False, ["Unusable answer tokens."]
    ans_grams = set(" ".join(ans_tokens[i:i+4]) for i in range(max(1, len(ans_tokens) - 3)))
    overlap = 0
    for cite in citations:
        doc_id = cite.get("doc_id")
        page = cite.get("page")
        # Prefer direct access to doc/page chunks:
        if hasattr(store, "get_doc_page_chunks"):
            candidates = store.get_doc_page_chunks(doc_id, page, limit=20)
        else:
            candidates = store.search(f"{doc_id} page {page}", k=20)
        for ch in candidates:
            c_tokens = re.findall(r"[A-Za-z0-9]+", (ch.text or "").lower())
            c_grams = set(" ".join(c_tokens[i:i+4]) for i in range(max(1, len(c_tokens) - 3)))
            if ans_grams & c_grams:
                overlap += 1
                break
    if overlap == 0:
        return False, ["Answer appears ungrounded in cited snippets."]
    return True, []

def validate_answer(question: str, answer: str) -> List[str]:
    """Lightweight completeness/consistency checks based on cues in the question."""
    issues = []
    ql = (question or "").lower()
    al = (answer or "").lower()

    # Require at least one numeral when SLAs/uptime/retention/etc. are asked
    if any(k in ql for k in ["sla", "uptime", "availability", "latency", "retention", "backup", "rpo", "rto"]):
        if not re.search(r"\d", al):
            issues.append("Missing numeric details for SLA/availability/retention.")

    # If question asks to "list/describe/provide", expect multi-sentence or bullets
    if any(k in ql for k in ["list", "describe", "provide", "explain", "outline"]):
        # crude heuristic: at least two sentences or two lines
        sentences = re.split(r"[.!?]\s+", normalize_text(answer))
        lines = [ln for ln in normalize_text(answer).splitlines() if ln.strip()]
        if len(sentences) < 2 and len(lines) < 2:
            issues.append("Response seems too brief for a 'list/describe/provide' prompt.")

    # If question mentions "encryption"/"security", expect key terms
    if any(k in ql for k in ["encryption", "crypto", "security"]):
        if not any(term in al for term in ["aes", "tls", "rsa", "sha"]):
            issues.append("Security answer missing common cryptographic terms (e.g., AES, TLS).")

    return issues

def compliance_gaps(question: str, answer: str) -> List[str]:
    """Rule-based compliance checks (env JSON)."""
    ql = (question or "").lower()
    al = (answer or "").lower()
    gaps = []
    for rule in COMPLIANCE_RULES:
        triggers = [t.lower() for t in rule.get("when_any", [])]
        if triggers and not any(t in ql or t in al for t in triggers):
            continue
        required = rule.get("must_include", [])
        missing = [req for req in required if req.lower() not in al]
        if missing:
            ctrl = rule.get("control", "policy")
            gaps.append(f"Missing required compliance terms {missing} for {ctrl}.")
    return gaps

# =========================
# DRAFTING (LLM)
# =========================
DRAFT_SYSTEM = (
    "You are the Drafting Agent. Use ONLY the provided context snippets to answer. "
    "Do not invent facts. Keep tone professional. "
    "Return ONLY JSON: {\"answer\":\"...\",\"status\":\"ok|needs_more_context\"}"
)

def extract_questions_llm(text: str) -> List[str]:
    prompt = f"""
Extract EVERY question from this RFP text. Keep multi-part items together.
Output ONLY a JSON array of strings.

Text:
{text[:120000]}
"""
    resp = client.chat.completions.create(
        model=EXTRACT_MODEL,
        temperature=0,
        messages=[
            {"role":"system","content":"You extract RFP questions precisely and output valid JSON."},
            {"role":"user","content":prompt}
        ]
    )
    raw = (resp.choices[0].message.content or "").strip().strip("`").strip()
    try:
        arr = json.loads(raw)
        return [q.strip() for q in arr if isinstance(q, str) and q.strip()]
    except Exception:
        # Fallback: naive tokenization on '?'
        parts = [p.strip()+"?" for p in text.split("?") if p.strip()]
        return parts[:25]

def draft_answer_llm(question: str, ctx_chunks: List[Dict[str,Any]]) -> Dict[str, Any]:
    context = "\n\n".join([f"[{i+1}] {c['text']}" for i,c in enumerate(ctx_chunks[:6])])
    prompt = f"Question: {question}\n\nContext:\n{context}"
    resp = client.chat.completions.create(
        model=DRAFT_MODEL,
        temperature=0.1,
        messages=[{"role":"system","content":DRAFT_SYSTEM},
                  {"role":"user","content":prompt}]
    )
    raw = (resp.choices[0].message.content or "").strip().strip("`").strip()
    try:
        obj = json.loads(raw)
    except Exception:
        obj = {"answer": raw, "status": "ok"}
    # Normalize/clean the answer safely
    obj["answer"] = normalize_text(coerce_str(obj.get("answer")))
    if obj.get("status") not in ("ok", "needs_more_context"):
        obj["status"] = "ok"
    return obj

def pick_citations(ctx_chunks: List[Dict[str,Any]], n: int = 2) -> List[Dict[str,Any]]:
    cites = []
    for c in ctx_chunks[:max(1, n)]:
        cites.append({
            "doc_id": c.get("doc_id"),
            "page": c.get("page"),
            "offset": c.get("offset")
        })
    return cites

def qa_answer(question: str, answer: str, citations: List[Dict[str,Any]]) -> Dict[str,Any]:
    """Validation (completeness), grounding, privacy, compliance + risk."""
    findings: List[str] = []
    status = "pass"

    # Grounding
    ok_g, g_reasons = check_grounding(answer, citations)
    if not ok_g:
        status = "fail"
        findings.extend(g_reasons)

    # Validation (completeness/basic consistency)
    v_issues = validate_answer(question, answer)
    if v_issues:
        status = "fail"
        findings.extend(v_issues)

    # Compliance rules
    c_issues = compliance_gaps(question, answer)
    if c_issues:
        status = "fail"
        findings.extend(c_issues)

    # Privacy redaction
    redacted, redactions = redact_answer(answer)

    # Risk score
    risk = 0
    if status == "fail": risk += 60
    if redactions: risk += 20

    fix = None
    if status == "fail":
        fix = ("Provide a concise, grounded answer using the cited snippets; "
               "include the specific numeric and compliance details requested; "
               "avoid introducing any non-cited claims.")

    return {
        "status": status,
        "reasons": findings,
        "fix": fix,
        "redactions": redactions,
        "controls": [],
        "risk": risk,
        "answer_effective": normalize_text(redacted),
    }

# =========================
# EXPORT
# =========================
def export_docx(rfp_id: str, qaed: List[Dict[str,Any]]) -> str:
    if not Document:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = Document()
    doc.add_heading(f"RFP Response – {rfp_id}", level=1)
    for i, item in enumerate(qaed, start=1):
        q = normalize_text(item["question"])
        a_qa   = normalize_text(coerce_str(item["qa"].get("answer_effective")))
        a_draft= normalize_text(coerce_str(item["draft"].get("answer")))
        a = a_qa or a_draft or ""
        doc.add_heading(f"Q{i}. {q}", level=2)
        for para in (a or "").split("\n\n"):
            doc.add_paragraph(para)
    out_path = os.path.join(EXPORT_DIR, f"{rfp_id}.docx")
    doc.save(out_path)
    return out_path

# =========================
# FLASK APP
# =========================
app = Flask(__name__, static_folder="static", template_folder="templates")
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 40 * 1024 * 1024  # 40 MB uploads

@app.before_request
def _guard():
    # Allow UI & health & ingest (for quick testing) without API key
    if request.path in ("/", "/health", "/ingest/upload") or request.path.startswith("/package/"):
        return
    if not api_guard_ok():
        return jsonify({"error": "unauthorized"}), 401

@app.get("/health")
def health():
    return jsonify({"ok": True, "vector_backend": VECTOR_BACKEND})

@app.get("/")
def index():
    return render_template("index.html")

@app.post("/ingest/upload")
def ingest_upload():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400
    filename = f.filename or f"upload-{uuid.uuid4()}"
    path = os.path.join(DATA_DIR, filename)
    f.save(path)

    # reference flag via query or form
    reference = (
        request.args.get("reference", "false").lower() in ("1", "true", "yes") or
        request.form.get("reference", "false").lower() in ("1", "true", "yes")
    )
    try:
        count = ingest_file(path, reference=reference)
    except Exception as e:
        return jsonify({"error": f"Ingestion failed: {e}"}), 500
    return jsonify({"doc_id": filename, "chunks": count, "reference": reference})

@app.post("/rfp/complete")
def rfp_complete():
    """
    Upload an RFP → extract questions → for each:
      - retrieve context (reference-first)
      - draft answer (LLM)  [normalized output]
      - QA (validation/grounding/privacy/compliance)
    → export DOCX + return preview JSON
    """
    f = request.files.get("rfp")
    if not f:
        return jsonify({"error": "No RFP uploaded"}), 400

    limit = int(request.form.get("limit", "20"))
    kctx  = int(request.form.get("k", "6"))

    # Save and read text
    tmpname = f"rfp-{uuid.uuid4()}{os.path.splitext(f.filename or '')[1]}"
    rfp_path = os.path.join(DATA_DIR, tmpname)
    f.save(rfp_path)

    if rfp_path.lower().endswith(".pdf"):
        pages = pdf_to_pages(rfp_path)
        rfp_text = "\n\n".join([t for _, t in pages])
    elif rfp_path.lower().endswith(".docx"):
        rfp_text = docx_to_text(rfp_path)
    else:
        rfp_text = open(rfp_path, "r", encoding="utf-8").read()

    # 1) Extract questions
    questions = extract_questions_llm(rfp_text)[:limit]
    results: List[Dict[str,Any]] = []

    # 2) Draft + QA per question
    for q in questions:
        ctx = store.search(q, k=kctx, filters={"reference_first": True})
        ctx_payload = [{
            "doc_id": c.doc_id,
            "page": c.page,
            "offset": list(c.offset),
            "text": c.text
        } for c in ctx]

        draft = draft_answer_llm(q, ctx_payload)
        ans = normalize_text(coerce_str(draft.get("answer")))
        status = draft.get("status", "ok")

        citations = pick_citations(ctx_payload, n=2)
        qa = qa_answer(q, ans, citations)

        # Optional: keep QA guidance visible; do not auto-overwrite the answer silently.
        results.append({
            "question": q,
            "citations": citations,
            "draft": {"status": status, "answer": ans},
            "qa": qa
        })

    # 3) Export DOCX
    rfp_id = os.path.splitext(os.path.basename(rfp_path))[0]
    try:
        out_path = export_docx(rfp_id, results)
    except Exception as e:
        # Still return answers even if export fails
        return jsonify({
            "rfp_id": rfp_id,
            "count_questions": len(questions),
            "answers": results,
            "download_uri": None,
            "export_error": str(e)
        })

    return jsonify({
        "rfp_id": rfp_id,
        "count_questions": len(questions),
        "answers": results,
        "download_uri": f"/package/download/{os.path.basename(out_path)}"
    })

@app.get("/package/download/<fname>")
def package_download(fname):
    path = os.path.join(EXPORT_DIR, fname)
    if not os.path.exists(path):
        return jsonify({"error": "not found"}), 404
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    debug_mode = os.getenv("FLASK_DEBUG", "1") in ("1", "true", "True")
    app.run(host="0.0.0.0", port=port, debug=debug_mode)
