# app.py — SWBC RFP Assistant (Config-driven)
# Ingest (reference boost) → Draft (LLM) → QA (grounding + privacy + compliance) → DOCX export

import os, io, re, uuid, json
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader

try:
    from docx import Document
except Exception:
    Document = None  # DOCX export disabled if python-docx isn't installed

# =========================
# ENV + CONFIG + CLIENT
# =========================
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY is required in environment")

client = OpenAI(api_key=OPENAI_API_KEY)

DATA_DIR = os.getenv("DATA_DIR", "./data")
EXPORT_DIR = os.getenv("EXPORT_DIR", "./exports")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

ACTION_API_KEY = os.getenv("API_KEY")  # optional API key for server endpoints

# ---- Load swbc_config.json (optional) ----
CONFIG_PATH = os.getenv("SWBC_CONFIG_PATH", "swbc_config.json")

def _deep_get(d: dict, path: str, default=None):
    cur = d
    for key in path.split("."):
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur

SWBC: Dict[str, Any] = {}
if os.path.exists(CONFIG_PATH):
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            SWBC = json.load(f)
    except Exception as e:
        print(f"[WARN] Failed to load {CONFIG_PATH}: {e}")

# ---- Config-driven knobs (env falls back) ----
EMBED_MODEL   = _deep_get(SWBC, "models.embed",   os.getenv("EMBED_MODEL",   "text-embedding-3-small"))
DRAFT_MODEL   = _deep_get(SWBC, "models.draft",   os.getenv("DRAFT_MODEL",   "gpt-4o-mini"))
EXTRACT_MODEL = _deep_get(SWBC, "models.extract", os.getenv("EXTRACT_MODEL", "gpt-4o-mini"))

REFERENCE_DOC_IDS = set(_deep_get(
    SWBC, "retrieval.reference_doc_ids",
    [x.strip() for x in os.getenv("REFERENCE_DOC_IDS", "RFQ_C2.pdf").split(",") if x.strip()]
))
REFERENCE_BOOST = float(_deep_get(SWBC, "retrieval.reference_boost", float(os.getenv("REFERENCE_BOOST", "1.25"))))

VECTOR_BACKEND = os.getenv("VECTOR_BACKEND", "local").lower()

CHUNK_LEN     = int(_deep_get(SWBC, "ingestion.chunk_len", 900))
CHUNK_OVERLAP = int(_deep_get(SWBC, "ingestion.chunk_overlap", 120))

# Prompts / tone
DRAFT_SYSTEM = _deep_get(SWBC, "prompts.drafting.system",
    "You are the Drafting Agent. Use ONLY the provided context snippets to answer. "
    "Do not invent facts. Keep tone professional. "
    'Return ONLY JSON: {"answer":"...","status":"ok|needs_more_context"}'
)
VOICE_GUIDE = _deep_get(SWBC, "prompts.voice", "")

EXTRACT_SYSTEM = _deep_get(SWBC, "prompts.extraction.system",
    "You extract all RFP questions precisely and output valid JSON."
)

QA_FIX_PROMPT = _deep_get(SWBC, "prompts.qa.fix",
    "Provide a concise, grounded answer using the cited snippets; include any required compliance terms."
)

# Privacy & compliance from config
EXTRA_PII_REGEXES: List[re.Pattern] = []
for pat in _deep_get(SWBC, "privacy.regexes", []):
    try:
        EXTRA_PII_REGEXES.append(re.compile(pat, re.IGNORECASE))
    except Exception as e:
        print(f"[WARN] Bad privacy regex in config: {pat} ({e})")

COMPLIANCE_RULES: List[Dict[str, Any]] = _deep_get(SWBC, "compliance.rules", [])

# Branding in export
BRAND_HEADER = _deep_get(SWBC, "brand.header", None)
BRAND_FOOTER = _deep_get(SWBC, "brand.footer", None)

# =========================
# UTILS
# =========================
def api_guard_ok():
    if ACTION_API_KEY and request.headers.get("X-API-KEY") != ACTION_API_KEY:
        return False
    return True

def l2norm(x: np.ndarray) -> np.ndarray:
    n = np.linalg.norm(x)
    return x / (n + 1e-12)

def cosine(a: np.ndarray, b: np.ndarray) -> float:
    return float(np.dot(l2norm(a), l2norm(b)))

def embed_text(text: str) -> List[float]:
    text = (text or "").replace("\x00", " ")
    text = re.sub(r"<img[^>]*>", "", text)
    text = text[:8000]  # token safety
    resp = client.embeddings.create(model=EMBED_MODEL, input=text)
    return resp.data[0].embedding

# =========================
# LOCAL VECTOR STORE
# =========================
CHUNK_FILE = os.path.join(DATA_DIR, "chunks.jsonl")

@dataclass
class Chunk:
    chunk_id: str
    doc_id: str
    page: int
    offset: Tuple[int, int]
    text: str
    meta: Dict[str, Any]
    vector: Optional[List[float]] = None
    score: Optional[float] = None

class VectorStore:
    def upsert(self, chunks: List[Chunk]) -> None: ...
    def search(self, query: str, k: int = 6, filters: Optional[Dict[str, Any]] = None) -> List[Chunk]: ...
    def has_doc(self, doc_id: str) -> bool: ...
    def get_doc_page_chunks(self, doc_id: str, page: Optional[int], limit: int = 20) -> List[Chunk]:
        return []

class LocalJSONStore(VectorStore):
    def upsert(self, chunks: List[Chunk]) -> None:
        with open(CHUNK_FILE, "a", encoding="utf-8") as f:
            for c in chunks:
                if c.vector is None:
                    c.vector = embed_text(c.text)
                f.write(json.dumps(asdict(c)) + "\n")

    def _iter_chunks(self):
        if not os.path.exists(CHUNK_FILE):
            return
        with open(CHUNK_FILE, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                yield Chunk(**json.loads(line))

    def has_doc(self, doc_id: str) -> bool:
        for c in self._iter_chunks():
            if c.doc_id == doc_id:
                return True
        return False

    def search(self, query: str, k: int = 6, filters: Optional[Dict[str, Any]] = None) -> List[Chunk]:
        qvec = np.array(embed_text(query), dtype=np.float32)
        prefer_reference = (filters or {}).get("reference_first", False)
        results: List[Chunk] = []
        for c in self._iter_chunks():
            cvec = np.array(c.vector, dtype=np.float32)
            s = cosine(qvec, cvec)
            if prefer_reference or c.doc_id in REFERENCE_DOC_IDS or c.meta.get("reference") is True:
                s *= REFERENCE_BOOST
            c.score = s
            results.append(c)
        results.sort(key=lambda x: x.score or 0.0, reverse=True)
        return results[:max(1, min(k, 50))]

    def get_doc_page_chunks(self, doc_id: str, page: Optional[int], limit: int = 20) -> List[Chunk]:
        out: List[Chunk] = []
        for c in self._iter_chunks():
            if c.doc_id != doc_id:
                continue
            if page is not None and c.page != page:
                continue
            out.append(c)
            if len(out) >= limit:
                break
        return out

# Stubs for other backends (future)
class QdrantStore(VectorStore): ...
class PgVectorStore(VectorStore): ...
class AtlasVectorStore(VectorStore): ...

def make_store() -> VectorStore:
    vb = VECTOR_BACKEND
    if vb == "qdrant":  return QdrantStore()
    if vb == "pgvector":return PgVectorStore()
    if vb == "atlas":   return AtlasVectorStore()
    return LocalJSONStore()

store: VectorStore = make_store()

# =========================
# INGESTION (memory-safe)
# =========================
def iter_chunks(txt: str, chunk_len: int = CHUNK_LEN, overlap: int = CHUNK_OVERLAP):
    n = len(txt)
    i = 0
    while i < n:
        j = min(n, i + chunk_len)
        yield (i, j, txt[i:j])
        if j >= n:
            break
        i = max(0, j - overlap)

def clean_page_text(t: str) -> str:
    t = (t or "").replace("\x00", " ")
    if len(t) > 600_000:  # clamp pathological PDF pages
        t = t[:600_000]
    return t

def pdf_to_pages(path: str) -> List[Tuple[int, str]]:
    reader = PdfReader(path)
    pages = []
    for p, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        pages.append((p, t))
    return pages

def docx_to_text(path: str) -> str:
    if not Document:
        raise RuntimeError("python-docx is required for DOCX ingestion. pip install python-docx")
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def ingest_file(filepath: str, reference: bool = False) -> int:
    """
    Streams chunks to the vector store in small batches (low RAM).
    """
    doc_id = os.path.basename(filepath)
    total_chunks = 0

    def _upsert_batch(batch: List[Chunk]):
        nonlocal total_chunks
        if not batch: return
        store.upsert(batch)
        total_chunks += len(batch)
        batch.clear()

    batch: List[Chunk] = []
    BATCH_SIZE = 16
    ref_flag = reference or (doc_id in REFERENCE_DOC_IDS)

    if filepath.lower().endswith(".pdf"):
        for page_idx, text in pdf_to_pages(filepath):
            text = clean_page_text(text)
            if not text.strip():
                continue
            for start, end, span in iter_chunks(text):
                batch.append(Chunk(
                    chunk_id=str(uuid.uuid4()),
                    doc_id=doc_id,
                    page=page_idx,
                    offset=(start, end),
                    text=span,
                    meta={"reference": ref_flag}
                ))
                if len(batch) >= BATCH_SIZE:
                    _upsert_batch(batch)

    elif filepath.lower().endswith(".docx"):
        text = clean_page_text(docx_to_text(filepath))
        for start, end, span in iter_chunks(text):
            batch.append(Chunk(
                chunk_id=str(uuid.uuid4()),
                doc_id=doc_id,
                page=0,
                offset=(start, end),
                text=span,
                meta={"reference": ref_flag}
            ))
            if len(batch) >= BATCH_SIZE:
                _upsert_batch(batch)

    else:
        with open(filepath, "r", encoding="utf-8") as f:
            text = clean_page_text(f.read())
        for start, end, span in iter_chunks(text):
            batch.append(Chunk(
                chunk_id=str(uuid.uuid4()),
                doc_id=doc_id,
                page=0,
                offset=(start, end),
                text=span,
                meta={"reference": ref_flag}
            ))
        _upsert_batch(batch)

    _upsert_batch(batch)  # flush
    return total_chunks

# =========================
# QA HELPERS (privacy + grounding + compliance)
# =========================
EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_RE = re.compile(r"\b(?:\+?\d{1,2}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}\b")
SSN_RE   = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
CC_RE    = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
APIKEY_RE= re.compile(r"\b(?:sk-[A-Za-z0-9]{20,}|AKIA[0-9A-Z]{16})\b")

def luhn_ok(s: str) -> bool:
    digits = [int(d) for d in re.sub(r"\D", "", s)]
    if len(digits) < 13: return False
    checksum = 0
    parity = len(digits) % 2
    for i, d in enumerate(digits):
        if i % 2 == parity:
            d *= 2
            if d > 9: d -= 9
        checksum += d
    return (checksum % 10) == 0

def redact_answer(ans: str):
    redactions = []
    def _re(regex, token):
        nonlocal ans
        for m in list(regex.finditer(ans)):
            redactions.append({"type": token.strip("<>"), "span": [m.start(), m.end()]})
        ans = regex.sub(token, ans)

    _re(EMAIL_RE, "<REDACTED_EMAIL>")
    _re(PHONE_RE, "<REDACTED_PHONE>")
    _re(SSN_RE, "<REDACTED_SSN>")
    for m in list(CC_RE.finditer(ans)):
        if luhn_ok(m.group(0)):
            redactions.append({"type":"CREDIT_CARD","span":[m.start(),m.end()]})
            ans = ans[:m.start()] + "<REDACTED_CC>" + ans[m.end():]
    _re(APIKEY_RE, "<REDACTED_SECRET>")

    for extra in EXTRA_PII_REGEXES:
        for m in list(extra.finditer(ans)):
            redactions.append({"type":"CUSTOM_PII","span":[m.start(),m.end()]})
        ans = extra.sub("<REDACTED_CUSTOM>", ans)

    return ans, redactions

def check_grounding(answer: str, citations: List[Dict[str, Any]]) -> Tuple[bool, List[str]]:
    if not answer.strip():
        return False, ["Empty answer."]
    if not citations:
        return False, ["No citations provided."]

    ans_tokens = re.findall(r"[A-Za-z0-9]+", answer.lower())
    if not ans_tokens:
        return False, ["Unusable answer tokens."]
    ans_grams = set(" ".join(ans_tokens[i:i+4]) for i in range(max(1, len(ans_tokens) - 3)))

    overlap = 0
    for cite in citations:
        doc_id = cite.get("doc_id")
        page = cite.get("page")
        candidates = store.get_doc_page_chunks(doc_id, page, limit=20)
        for ch in candidates:
            c_tokens = re.findall(r"[A-Za-z0-9]+", (ch.text or "").lower())
            c_grams = set(" ".join(c_tokens[i:i+4]) for i in range(max(1, len(c_tokens) - 3)))
            if ans_grams & c_grams:
                overlap += 1
                break
    if overlap == 0:
        return False, ["Answer appears ungrounded in cited snippets."]
    return True, []

def compliance_gaps(question: str, answer: str) -> List[str]:
    ql = (question or "").lower()
    al = (answer or "").lower()
    gaps = []
    for rule in COMPLIANCE_RULES:
        triggers = [t.lower() for t in rule.get("when_any", [])]
        if triggers and not any(t in ql or t in al for t in triggers):
            continue
        required = rule.get("must_include", [])
        missing = [req for req in required if req.lower() not in al]
        if missing:
            ctrl = rule.get("control", "policy")
            gaps.append(f"Missing required compliance terms {missing} for {ctrl}.")
    return gaps

def qa_answer(answer: str, citations: List[Dict[str, Any]], question: str = "") -> Dict[str,Any]:
    status = "pass"
    reasons: List[str] = []

    ok_g, g_reasons = check_grounding(answer, citations)
    if not ok_g:
        status = "fail"; reasons.extend(g_reasons)

    comp = compliance_gaps(question, answer)
    if comp:
        status = "fail"; reasons.extend(comp)

    redacted, redactions = redact_answer(answer)
    risk = 0
    if status == "fail": risk += 60
    if redactions: risk += 20

    fix = QA_FIX_PROMPT if status == "fail" else None

    return {
        "status": status,
        "reasons": reasons,
        "fix": fix,
        "redactions": redactions,
        "controls": [],
        "risk": risk,
        "answer_effective": redacted
    }

# =========================
# DRAFTING (LLM)
# =========================
def extract_questions_llm(text: str) -> List[str]:
    prompt = f"""
Extract EVERY question from this RFP text. Keep multi-part items together.
Output ONLY a JSON array of strings.

Text:
{text[:120000]}
"""
    resp = client.chat.completions.create(
        model=EXTRACT_MODEL,
        temperature=0,
        messages=[
            {"role":"system","content":EXTRACT_SYSTEM},
            {"role":"user","content":prompt}
        ]
    )
    raw = (resp.choices[0].message.content or "").strip().strip("`").strip()
    try:
        arr = json.loads(raw)
        return [q.strip() for q in arr if isinstance(q, str) and q.strip()]
    except Exception:
        parts = [p.strip()+"?" for p in text.split("?") if p.strip()]
        return parts[:25]

def draft_answer_llm(question: str, ctx_chunks: List[Dict[str,Any]]) -> Dict[str, Any]:
    context = "\n\n".join([f"[{i+1}] {c['text']}" for i,c in enumerate(ctx_chunks[:6])])
    sys = DRAFT_SYSTEM + (("\n\nVoice:\n" + VOICE_GUIDE) if VOICE_GUIDE else "")
    prompt = f"Question: {question}\n\nContext:\n{context}"
    resp = client.chat.completions.create(
        model=DRAFT_MODEL,
        temperature=0.1,
        messages=[{"role":"system","content":sys},
                  {"role":"user","content":prompt}]
    )
    raw = (resp.choices[0].message.content or "").strip().strip("`").strip()
    try:
        return json.loads(raw)
    except Exception:
        return {"answer": raw, "status": "ok"}

def pick_citations(ctx_chunks: List[Dict[str,Any]], n: int = 2) -> List[Dict[str,Any]]:
    cites = []
    for c in ctx_chunks[:max(1, n)]:
        cites.append({
            "doc_id": c.get("doc_id"),
            "page": c.get("page"),
            "offset": c.get("offset")
        })
    return cites

# =========================
# EXPORT
# =========================
def export_docx(rfp_id: str, qaed: List[Dict[str,Any]]) -> str:
    if not Document:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = Document()
    title = _deep_get(SWBC, "brand.title", f"RFP Response – {rfp_id}")
    doc.add_heading(title, level=1)

    if BRAND_HEADER:
        doc.add_paragraph(BRAND_HEADER)

    for i, item in enumerate(qaed, start=1):
        q = item["question"]
        a = item["qa"]["answer_effective"] or item["draft"]["answer"] or ""
        doc.add_heading(f"Q{i}. {q}", level=2)
        for para in (a or "").split("\n\n"):
            doc.add_paragraph(para)

    if BRAND_FOOTER:
        doc.add_paragraph("\n" + BRAND_FOOTER)

    out_path = os.path.join(EXPORT_DIR, f"{rfp_id}.docx")
    doc.save(out_path)
    return out_path

# =========================
# FLASK APP
# =========================
app = Flask(__name__, static_folder="static", template_folder="templates")
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 40 * 1024 * 1024  # 40 MB uploads

@app.before_request
def _guard():
    # Allow UI, health, ingest, and download without API key for local testing
    if request.path in ("/", "/health", "/ingest/upload") or request.path.startswith("/package/"):
        return
    if not api_guard_ok():
        return jsonify({"error": "unauthorized"}), 401

@app.get("/health")
def health():
    return jsonify({
        "ok": True,
        "vector_backend": VECTOR_BACKEND,
        "models": {"embed": EMBED_MODEL, "draft": DRAFT_MODEL, "extract": EXTRACT_MODEL},
        "reference_boost": REFERENCE_BOOST,
        "reference_doc_ids": sorted(list(REFERENCE_DOC_IDS)),
        "chunking": {"len": CHUNK_LEN, "overlap": CHUNK_OVERLAP}
    })

@app.get("/")
def index():
    return render_template("index.html")

@app.post("/ingest/upload")
def ingest_upload():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file"}), 400
    filename = f.filename or f"upload-{uuid.uuid4()}"
    path = os.path.join(DATA_DIR, filename)
    f.save(path)

    reference = (
        request.args.get("reference", "false").lower() in ("1", "true", "yes") or
        request.form.get("reference", "false").lower() in ("1", "true", "yes")
    )
    try:
        count = ingest_file(path, reference=reference)
    except Exception as e:
        return jsonify({"error": f"Ingestion failed: {e}"}), 500
    return jsonify({"doc_id": filename, "chunks": count, "reference": reference})

@app.post("/rfp/complete")
def rfp_complete():
    """
    Upload an RFP → extract questions → for each:
      - retrieve context (reference-first)
      - draft answer (LLM)
      - QA (validation/grounding/privacy/compliance)
    → export DOCX + return preview JSON
    """
    f = request.files.get("rfp")
    if not f:
        return jsonify({"error": "No RFP uploaded"}), 400

    limit = int(request.form.get("limit", "20"))
    kctx  = int(request.form.get("k", "6"))

    tmpname = f"rfp-{uuid.uuid4()}{os.path.splitext(f.filename or '')[1]}"
    rfp_path = os.path.join(DATA_DIR, tmpname)
    f.save(rfp_path)

    if rfp_path.lower().endswith(".pdf"):
        pages = pdf_to_pages(rfp_path)
        rfp_text = "\n\n".join([t for _, t in pages])
    elif rfp_path.lower().endswith(".docx"):
        rfp_text = docx_to_text(rfp_path)
    else:
        rfp_text = open(rfp_path, "r", encoding="utf-8").read()

    # 1) Extract questions
    questions = extract_questions_llm(rfp_text)[:limit]
    results: List[Dict[str,Any]] = []

    # 2) Draft + QA per question
    for q in questions:
        ctx = store.search(q, k=kctx, filters={"reference_first": True})
        ctx_payload = [{
            "doc_id": c.doc_id,
            "page": c.page,
            "offset": list(c.offset),
            "text": c.text
        } for c in ctx]

        draft = draft_answer_llm(q, ctx_payload)
        ans = (draft.get("answer") or "").strip()
        status = draft.get("status", "ok")

        citations = pick_citations(ctx_payload, n=2)
        qa = qa_answer(ans, citations, question=q)

        # Optional one-shot fix if fail: ask again with fix guidance + same citations
        if qa["status"] == "fail" and qa.get("fix"):
            # You can re-run the drafter with a stricter instruction here if you want.
            # For now, just keep QA's guidance visible in UI.
            pass

        results.append({
            "question": q,
            "citations": citations,
            "draft": {"status": status, "answer": ans},
            "qa": qa
        })

    # 3) Export DOCX
    rfp_id = os.path.splitext(os.path.basename(rfp_path))[0]
    try:
        out_path = export_docx(rfp_id, results)
    except Exception as e:
        return jsonify({
            "rfp_id": rfp_id,
            "count_questions": len(questions),
            "answers": results,
            "download_uri": None,
            "export_error": str(e)
        })

    return jsonify({
        "rfp_id": rfp_id,
        "count_questions": len(questions),
        "answers": results,
        "download_uri": f"/package/download/{os.path.basename(out_path)}"
    })

@app.get("/package/download/<fname>")
def package_download(fname):
    path = os.path.join(EXPORT_DIR, fname)
    if not os.path.exists(path):
        return jsonify({"error": "not found"}), 404
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    debug_mode = os.getenv("FLASK_DEBUG", "1") in ("1", "true", "True")
    app.run(host="0.0.0.0", port=port, debug=debug_mode)
